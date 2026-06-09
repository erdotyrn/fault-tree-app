#!/usr/bin/env python3
"""HÖA Analizörü Proje Raporu — Word (.docx) oluşturucu."""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shd)

def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar1)
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run._element.append(instrText)
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar2)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(120, 120, 120)

def add_img(doc, img_name, caption):
    ss = os.path.join(os.path.dirname(__file__), "screenshots", img_name)
    if not os.path.exists(ss):
        return
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(ss, width=Inches(5.8))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

def create_report():
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ══════════════════════════════════════
    #  KAPAK
    # ══════════════════════════════════════

    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("HÖA ANALİZÖRÜ")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(25, 118, 210)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Hata Ağacı & Olay Ağacı Analiz Uygulaması")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    info_lines = [
        ("Hazırlayan", "Erdoğan Toyran"),
        ("Bölüm", "Nükleer Enerji Mühendisliği"),
        ("Ders", "Risk Analizleri"),
        ("Danışman", "Hüseyin Sahiner"),
        ("Tarih", "Haziran 2026"),
    ]
    table = doc.add_table(rows=len(info_lines), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(info_lines):
        c0 = table.cell(i, 0)
        c1 = table.cell(i, 1)
        c0.width = Cm(4)
        c1.width = Cm(8)
        r0 = c0.paragraphs[0].add_run(f"{label}:")
        r0.bold = True
        r0.font.size = Pt(12)
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r1 = c1.paragraphs[0].add_run(f"  {value}")
        r1.font.size = Pt(12)

    doc.add_page_break()

    # ══════════════════════════════════════
    #  ÖZET
    # ══════════════════════════════════════

    h = doc.add_heading("Özet", level=1)
    h.runs[0].font.color.rgb = RGBColor(25, 118, 210)

    doc.add_paragraph(
        "Bu projede Hata Ağacı Analizi (FTA), Olay Ağacı Analizi (ETA) "
        "ve Seri/Paralel Sistem Analizi yöntemlerini tek bir masaüstü "
        "uygulamasında birleştirdim. Python ve PyQt6 ile geliştirdiğim "
        "uygulama; AND, OR, VOTE(k/n) mantık geçitleri, minimal kesim "
        "kümeleri, önem ölçüleri, zamana bağlı güvenilirlik hesabı, "
        "otomatik diyagram çizimi ve PDF rapor üretimi yapabiliyor. "
        "Yaklaşık 4000 satır koddan oluşan uygulama, nükleer güvenlik "
        "başta olmak üzere risk analizi eğitiminde kullanılabilecek "
        "bir araç olarak tasarlandı."
    )

    p = doc.add_paragraph()
    run = p.add_run("Anahtar Kelimeler: ")
    run.bold = True
    p.add_run(
        "Hata Ağacı Analizi, Olay Ağacı Analizi, Güvenilirlik, "
        "Minimal Kesim Kümeleri, Önem Ölçüleri"
    )

    doc.add_page_break()

    # ══════════════════════════════════════
    #  SEMBOL LİSTESİ
    # ══════════════════════════════════════

    h = doc.add_heading("Sembol Listesi", level=1)
    h.runs[0].font.color.rgb = RGBColor(25, 118, 210)

    symbols = [
        ("R", "Güvenilirlik — bileşenin çalışma olasılığı (0 ile 1 arası)"),
        ("F", "Arıza olasılığı (F = 1 − R)"),
        ("λ", "Arıza oranı (1/saat veya 1/yıl)"),
        ("t", "Görev süresi (saat veya yıl)"),
        ("P(tepe)", "Tepe olayın gerçekleşme olasılığı"),
        ("MCS", "Minimal Kesim Kümesi"),
        ("k/n", "n elemandan k tanesinin arızalanması gereken oylama sistemi"),
        ("f", "Başlatıcı olay frekansı (olay/yıl)"),
    ]

    sym_table = doc.add_table(rows=len(symbols) + 1, cols=2)
    sym_table.style = 'Table Grid'
    sym_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h_text in enumerate(["Sembol", "Tanım"]):
        cell = sym_table.cell(0, i)
        cell.paragraphs[0].add_run(h_text).bold = True
        set_cell_shading(cell, "1976D2")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, (sym, desc) in enumerate(symbols, 1):
        r = sym_table.cell(row_idx, 0).paragraphs[0].add_run(sym)
        r.font.name = 'Cambria Math'
        r.bold = True
        sym_table.cell(row_idx, 1).text = desc
        if row_idx % 2 == 0:
            for ci in range(2):
                set_cell_shading(sym_table.cell(row_idx, ci), "E3F2FD")

    doc.add_page_break()

    # ══════════════════════════════════════
    #  İÇİNDEKİLER
    # ══════════════════════════════════════

    h = doc.add_heading("İÇİNDEKİLER", level=1)
    h.runs[0].font.color.rgb = RGBColor(25, 118, 210)

    toc_items = [
        "1. Giriş ve Amaç",
        "2. Teorik Arka Plan",
        "   2.1. Hata Ağacı Analizi (FTA)",
        "   2.2. Minimal Kesim Kümeleri (MCS)",
        "   2.3. Önem Ölçüleri",
        "   2.4. Olay Ağacı Analizi (ETA)",
        "   2.5. Seri-Paralel Sistem Analizi",
        "   2.6. Zamana Bağlı Model",
        "   2.7. Varsayımlar ve Sınırlamalar",
        "3. Yazılım Tasarımı",
        "4. Uygulama Özellikleri",
        "5. Kullanım Rehberi",
        "6. Örnek Analizler",
        "7. Doğrulama",
        "8. Sonuç",
        "Kaynaklar",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        if item.startswith("   "):
            p.paragraph_format.left_indent = Cm(1.5)

    doc.add_page_break()

    # ══════════════════════════════════════
    #  1. GİRİŞ
    # ══════════════════════════════════════

    h = doc.add_heading("1. Giriş ve Amaç", level=1)
    h.runs[0].font.color.rgb = RGBColor(25, 118, 210)

    doc.add_paragraph(
        "Nükleer santraller, kimya tesisleri, havacılık gibi sektörlerde "
        "bir arıza ciddi sonuçlar doğurabilir. Bu yüzden tasarım aşamasında "
        "\"bu sistem ne kadar güvenilir?\" ve \"en zayıf halkası neresi?\" "
        "sorularını sormamız gerekir. Güvenilirlik mühendisliği tam olarak "
        "bununla ilgilenir."
    )

    doc.add_paragraph(
        "Derste gördüğümüz üç temel yöntem var: Hata Ağacı Analizi (FTA) ile "
        "bir arızanın kök nedenlerini buluyoruz, Olay Ağacı Analizi (ETA) ile "
        "bir kaza başladıktan sonra olası senaryoları görüyoruz, Seri/Paralel "
        "analizle de sistemin genel güvenilirliğini hesaplıyoruz."
    )

    doc.add_paragraph(
        "Bu projede bu üç yöntemi tek bir masaüstü uygulamasında birleştirdim. "
        "Kullanıcı olayları ve mantık yapılarını tanımlıyor, program otomatik "
        "olarak hesaplıyor, diyagram çiziyor ve PDF rapor üretiyor. Python "
        "dili ve PyQt6 arayüz kütüphanesi kullandım."
    )

    # ══════════════════════════════════════
    #  2. TEORİK ARKA PLAN
    # ══════════════════════════════════════

    h = doc.add_heading("2. Teorik Arka Plan", level=1)
    h.runs[0].font.color.rgb = RGBColor(25, 118, 210)

    # --- 2.1 FTA ---
    h2 = doc.add_heading("2.1. Hata Ağacı Analizi (FTA)", level=2)
    h2.runs[0].font.color.rgb = RGBColor(230, 81, 0)

    doc.add_paragraph(
        "Hata ağacında en üste istemediğimiz bir olay koyuyoruz — buna "
        "\"tepe olay\" diyoruz, mesela \"soğutma kaybı\". Sonra bu olayın "
        "hangi durumlarda gerçekleşebileceğini aşağıya doğru açıyoruz. "
        "En alttaki yapraklar temel olaylar (pompa arızası, vana sıkışması "
        "gibi), aradaki düğümler ise mantık geçitleri."
    )

    doc.add_paragraph("Üç tip mantık geçidi kullanıyoruz:")

    # AND
    p = doc.add_paragraph()
    run = p.add_run("AND (VE) geçidi: ")
    run.bold = True
    p.add_run(
        "Alt elemanların hepsinin aynı anda arızalanması gerekir. "
        "Yedekli sistemleri modellemek için kullanırız. Mesela iki paralel "
        "pompanın ikisinin de durması lazım — biri çalışıyorsa sistem devam eder."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("R(geçit) = 1 − ∏(1 − Rᵢ)")
    run.bold = True
    run.font.color.rgb = RGBColor(21, 101, 192)

    doc.add_paragraph(
        "Yani her bir elemanın arıza olasılıklarını çarpıp 1'den çıkarıyoruz. "
        "Örnek: Pompa A (R=0.99) ve Pompa B (R=0.98) paralelse → "
        "R = 1 − (0.01 × 0.02) = 0.9998. Yedeklilik sayesinde güvenilirlik "
        "çok yükseldi."
    )

    # OR
    p = doc.add_paragraph()
    run = p.add_run("OR (VEYA) geçidi: ")
    run.bold = True
    p.add_run(
        "Alt elemanlardan herhangi birinin arızalanması yeterli. "
        "Tek bir arızanın sistemi durdurduğu durumlarda kullanırız."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("R(geçit) = ∏ Rᵢ")
    run.bold = True
    run.font.color.rgb = RGBColor(21, 101, 192)

    doc.add_paragraph(
        "Tüm elemanların güvenilirliklerini çarpıyoruz. "
        "Örnek: İki arıza modu (R₁=0.99 ve R₂=0.98) varsa → "
        "R = 0.99 × 0.98 = 0.9702. Herhangi biri olunca sistem çöktüğü "
        "için güvenilirlik düşüyor."
    )

    # VOTE
    p = doc.add_paragraph()
    run = p.add_run("VOTE (k/n) geçidi: ")
    run.bold = True
    p.add_run(
        "n elemandan en az k tanesinin arızalanması gerekir. AND ve OR'un "
        "genelleştirilmiş hali. Nükleer santrallerdeki 2/3 oylama "
        "sistemlerinde sıkça karşımıza çıkar. k = n ise AND, k = 1 ise "
        "OR geçidine eşdeğer olur."
    )

    doc.add_paragraph(
        "Burada R güvenilirlik, yani bileşenin çalışma olasılığını ifade "
        "ediyor. Arıza olasılığı ise basitçe F = 1 − R."
    )

    # --- 2.2 MCS ---
    h2 = doc.add_heading("2.2. Minimal Kesim Kümeleri (MCS)", level=2)
    h2.runs[0].font.color.rgb = RGBColor(230, 81, 0)

    doc.add_paragraph(
        "Minimal kesim kümesi, tepe olayı tek başına gerçekleştirebilen en "
        "küçük olay grubu. Örneğin {Pompa A, Pompa B} bir MCS ise bu ikisi "
        "birlikte arızalanınca sistem çöker, ama tek başına biri yetmez."
    )

    doc.add_paragraph(
        "MCS'leri bulmak için MOCUS algoritmasını kullandım. Algoritma tepe "
        "olaydan başlayarak geçitleri açar: OR geçidi alternatif yollar "
        "üretir, AND geçidi eşzamanlı gereksinimleri birleştirir. "
        "Sonunda elde edilen kümeler minimize edilir."
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("P(tepe) ≈ Σ P(MCSᵢ)   (Nadir olay yaklaşımı)")
    run.bold = True
    run.font.color.rgb = RGBColor(21, 101, 192)

    doc.add_paragraph(
        "Yani tepe olayın olasılığını bulmak için her MCS'nin olasılığını "
        "toplayabiliriz. Bu yaklaşım olasılıklar düşükken iyi çalışır."
    )

    # --- 2.3 Önem Ölçüleri ---
    h2 = doc.add_heading("2.3. Önem Ölçüleri", level=2)
    h2.runs[0].font.color.rgb = RGBColor(230, 81, 0)

    doc.add_paragraph(
        "Hangi bileşenin en kritik olduğunu bulmak için önem ölçülerini "
        "hesaplıyoruz. Temel soru: \"Hangi bileşeni iyileştirirsem sistemi "
        "en çok güçlendiririm?\""
    )

    imp_table = doc.add_table(rows=5, cols=3)
    imp_table.style = 'Table Grid'
    imp_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Ölçü", "Ne hesaplıyor?", "Nasıl yorumlarız?"]
    for i, h_text in enumerate(headers):
        cell = imp_table.cell(0, i)
        cell.paragraphs[0].add_run(h_text).bold = True
        set_cell_shading(cell, "1976D2")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    imp_data = [
        ("Birnbaum",
         "P(tepe|qᵢ=1) − P(tepe|qᵢ=0)",
         "Bileşenin tepe olaya marjinal etkisi"),
        ("Fussell-Vesely",
         "(P(tepe) − P(tepe|qᵢ=0)) / P(tepe)",
         "Toplam riske yüzde kaç katkı yapıyor"),
        ("RAW",
         "P(tepe|qᵢ=1) / P(tepe)",
         "Bu bileşen bozulursa risk kaç kat artar"),
        ("RRW",
         "P(tepe) / P(tepe|qᵢ=0)",
         "Bu bileşen mükemmel olsa risk kaç kat düşer"),
    ]
    for row_idx, (name, formula, desc) in enumerate(imp_data, 1):
        imp_table.cell(row_idx, 0).text = name
        imp_table.cell(row_idx, 1).text = formula
        imp_table.cell(row_idx, 2).text = desc
        if row_idx % 2 == 0:
            for ci in range(3):
                set_cell_shading(imp_table.cell(row_idx, ci), "E3F2FD")

    doc.add_paragraph()

    doc.add_paragraph(
        "Burada qᵢ = 1 demek \"bu bileşen kesinlikle bozuk\", "
        "qᵢ = 0 demek \"bu bileşen kesinlikle sağlam\" demek. "
        "Mesela RAW = 10 ise o bileşen bozulduğunda risk 10 kat artıyor."
    )

    # --- 2.4 ETA ---
    h2 = doc.add_heading("2.4. Olay Ağacı Analizi (ETA)", level=2)
    h2.runs[0].font.color.rgb = RGBColor(230, 81, 0)

    doc.add_paragraph(
        "Olay ağacı, FTA'nın tersine çalışır. Bir başlatıcı olaydan "
        "(mesela boru kırılması) ileriye doğru bakar. Sırayla güvenlik "
        "sistemleri devreye girer, her biri ya başarılı olur ya başarısız. "
        "Bu dallanmalar bize olası senaryoları verir."
    )

    doc.add_paragraph(
        "N tane güvenlik bariyeri varsa 2ᴺ farklı sonuç elde ederiz. "
        "Her sonucun olasılığını dallar boyunca çarparak buluruz:"
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("P(sonuç) = ∏ P(dalᵢ)")
    run.bold = True
    run.font.color.rgb = RGBColor(21, 101, 192)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Frekans = f(başlatıcı olay) × P(sonuç)")
    run.bold = True
    run.font.color.rgb = RGBColor(21, 101, 192)

    doc.add_paragraph(
        "Mesela başlatıcı olay frekansı 10⁻⁴/yıl ve bir senaryonun "
        "olasılığı 0.01 ise o senaryonun frekansı 10⁻⁶/yıl olur."
    )

    # --- 2.5 Seri/Paralel ---
    h2 = doc.add_heading("2.5. Seri-Paralel Sistem Analizi", level=2)
    h2.runs[0].font.color.rgb = RGBColor(230, 81, 0)

    doc.add_paragraph(
        "Bileşenlerin nasıl bağlandığına göre sistem güvenilirliği değişir:"
    )

    sp_table = doc.add_table(rows=3, cols=3)
    sp_table.style = 'Table Grid'
    sp_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    sp_headers = ["Bağlantı", "Formül", "Ne demek?"]
    for i, h_text in enumerate(sp_headers):
        cell = sp_table.cell(0, i)
        cell.paragraphs[0].add_run(h_text).bold = True
        set_cell_shading(cell, "1976D2")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    sp_data = [
        ("Seri", "R = R₁ × R₂ × ... × Rₙ",
         "Biri bozulursa sistem durur"),
        ("Paralel", "R = 1 − (1−R₁)(1−R₂)...(1−Rₙ)",
         "Hepsi bozulursa ancak o zaman sistem durur"),
    ]
    for row_idx, (name, formula, desc) in enumerate(sp_data, 1):
        sp_table.cell(row_idx, 0).text = name
        sp_table.cell(row_idx, 1).text = formula
        sp_table.cell(row_idx, 2).text = desc

    doc.add_paragraph()

    doc.add_paragraph(
        "Dikkat edilirse: seri bağlantıda güvenilirlikler çarpılıyor "
        "(OR geçidiyle aynı mantık — herhangi biri bozulursa yeter), "
        "paralel bağlantıda ise arıza olasılıkları çarpılıp 1'den "
        "çıkarılıyor (AND geçidiyle aynı mantık — hepsi bozulmalı). "
        "Yani seri = zayıf halka, paralel = yedeklilik."
    )

    # --- 2.6 Zamana Bağlı ---
    h2 = doc.add_heading("2.6. Zamana Bağlı Model", level=2)
    h2.runs[0].font.color.rgb = RGBColor(230, 81, 0)

    doc.add_paragraph(
        "Bazen güvenilirliği doğrudan R olarak vermek yerine arıza oranı "
        "(λ) ve görev süresi (t) ile hesaplamak isteriz:"
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("R(t) = e^(−λt)")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(21, 101, 192)

    doc.add_paragraph(
        "Mesela bir pompanın arıza oranı λ = 0.0003/saat ve görev süresi "
        "t = 8760 saat (1 yıl) ise: R = e^(−0.0003 × 8760) ≈ 0.073. "
        "Yani yıl boyunca kesintisiz çalışma olasılığı %7.3."
    )

    # --- 2.7 Varsayımlar ve Sınırlamalar ---
    h2 = doc.add_heading("2.7. Varsayımlar ve Sınırlamalar", level=2)
    h2.runs[0].font.color.rgb = RGBColor(230, 81, 0)

    doc.add_paragraph(
        "Uygulamadaki hesaplamalar şu varsayımlara dayanıyor:"
    )

    limits = [
        ("Bağımsızlık",
         "Tüm temel olaylar birbirinden bağımsız kabul ediliyor. "
         "Gerçekte ortak neden arızaları (mesela aynı deprem iki pompayı "
         "birden bozabilir) bu varsayımı bozabilir."),
        ("Nadir olay yaklaşımı",
         "MCS olasılık toplamında olasılıklar düşükken "
         "iyi çalışır, yüksek olasılıklarda hata payı artar."),
        ("Sadece üstel dağılım",
         "Zamana bağlı modelde yalnızca R = e^(−λt) var. "
         "Weibull gibi yaşlanma modelleri henüz desteklenmiyor."),
        ("İkili dallanma",
         "Olay ağacında her dal sadece başarı/başarısız olarak ayrılıyor."),
    ]

    for title, desc in limits:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}: ")
        run.bold = True
        p.add_run(desc)

    # ══════════════════════════════════════
    #  3. YAZILIM TASARIMI
    # ══════════════════════════════════════

    h = doc.add_heading("3. Yazılım Tasarımı", level=1)
    h.runs[0].font.color.rgb = RGBColor(25, 118, 210)

    doc.add_paragraph(
        "Uygulamayı Python ile yazdım. Arayüz için PyQt6, diyagramlar için "
        "Graphviz, PDF çıktısı için fpdf2 kütüphanelerini kullandım."
    )

    tech_table = doc.add_table(rows=5, cols=3)
    tech_table.style = 'Table Grid'
    tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tech_headers = ["Teknoloji", "Sürüm", "Görevi"]
    for i, h_text in enumerate(tech_headers):
        cell = tech_table.cell(0, i)
        cell.paragraphs[0].add_run(h_text).bold = True
        set_cell_shading(cell, "1976D2")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    tech_data = [
        ("Python", "3.10+", "Ana programlama dili"),
        ("PyQt6", "6.5+", "Grafik arayüz"),
        ("Graphviz", "0.20+", "Ağaç diyagramları"),
        ("fpdf2", "2.7+", "PDF rapor çıktısı"),
    ]
    for row_idx, (tech, ver, purpose) in enumerate(tech_data, 1):
        tech_table.cell(row_idx, 0).text = tech
        tech_table.cell(row_idx, 1).text = ver
        tech_table.cell(row_idx, 2).text = purpose

    doc.add_paragraph()

    doc.add_paragraph("Kodu dört modüle ayırdım:")

    mod_table = doc.add_table(rows=5, cols=3)
    mod_table.style = 'Table Grid'
    mod_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    mod_headers = ["Dosya", "Satır", "Ne yapıyor?"]
    for i, h_text in enumerate(mod_headers):
        cell = mod_table.cell(0, i)
        cell.paragraphs[0].add_run(h_text).bold = True
        set_cell_shading(cell, "1976D2")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    mod_data = [
        ("models.py", "~330",
         "Veri yapıları — olay, geçit, ağaç, proje tanımları ve JSON kaydet/yükle"),
        ("engine.py", "~415",
         "Hesaplama motorları — FTA, ETA, MCS ve önem ölçüleri"),
        ("visualization.py", "~550",
         "Diyagram çizimi — Graphviz ve QPainter ile geçit sembolleri"),
        ("app.py", "~2700",
         "Arayüz — üç mod, diyaloglar, örnekler, PDF rapor, kılavuz"),
    ]
    for row_idx, (mod, lines, desc) in enumerate(mod_data, 1):
        r = mod_table.cell(row_idx, 0).paragraphs[0].add_run(mod)
        r.bold = True
        r.font.name = 'Consolas'
        mod_table.cell(row_idx, 1).text = lines
        mod_table.cell(row_idx, 2).text = desc

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Toplam: ")
    run.bold = True
    p.add_run("yaklaşık 4000 satır kod.")

    doc.add_paragraph()

    doc.add_paragraph(
        "Veri akışı şu şekilde: Kullanıcı arayüzde olay ve geçitleri "
        "tanımlıyor → models.py'ye kaydediliyor → engine.py hesaplıyor → "
        "visualization.py diyagram çiziyor → sonuçlar ekrana yansıyor. "
        "Her değişiklikte bu zincir otomatik çalışıyor."
    )

    # ══════════════════════════════════════
    #  4. ÖZELLİKLER
    # ══════════════════════════════════════

    h = doc.add_heading("4. Uygulama Özellikleri", level=1)
    h.runs[0].font.color.rgb = RGBColor(25, 118, 210)

    features = [
        ("Üç analiz modu",
         "FTA, ETA ve Seri/Paralel tek pencerede, üstteki butonlarla geçiş yapılıyor."),
        ("AND / OR / VOTE(k/n) mantık geçitleri",
         "Standart VE/VEYA'ya ek olarak, nükleer santrallerde kullanılan "
         "k/n oylama geçidi de var."),
        ("Minimal Kesim Kümeleri",
         "MOCUS algoritması ile sistemi çökerten en küçük bileşen "
         "gruplarını otomatik buluyor."),
        ("Önem ölçüleri",
         "Birnbaum, Fussell-Vesely, RAW, RRW — hangi bileşen en kritik, hemen görünüyor."),
        ("Otomatik hesaplama",
         "Her değişiklikte sonuçlar ve diyagram anında güncelleniyor, "
         "ayrıca bir \"hesapla\" butonuna basmaya gerek yok."),
        ("Diyagramlar",
         "Graphviz ile standart mühendislik notasyonunda renkli ağaç diyagramları. "
         "AND/OR sembolleri QPainter ile çiziliyor."),
        ("Zamana bağlı model",
         "R = e^(−λt) hesabı doğrudan olay tanımlama ekranında yapılabiliyor."),
        ("PDF rapor",
         "Tek tuşla tüm sonuçları ve diyagramı PDF olarak dışa aktarma."),
        ("Kaydet / Yükle",
         "Projeler JSON formatında saklanıyor, istediğin zaman geri yükleyebilirsin."),
        ("8 hazır örnek",
         "Nükleer, kimyasal, endüstriyel — her analiz türü için açıklamalı örnek projeler."),
        ("Kullanım kılavuzu",
         "İlk açılışta hoş geldiniz ekranı ve Yardım menüsünde adım adım kılavuz."),
        ("Diyagram kaydetme",
         "Oluşan diyagramı PNG olarak dosyaya kaydedebilme."),
    ]

    for i, (title_text, desc) in enumerate(features, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {title_text}: ")
        run.bold = True
        p.add_run(desc)

    # ══════════════════════════════════════
    #  5. KULLANIM REHBERİ
    # ══════════════════════════════════════

    h = doc.add_heading("5. Kullanım Rehberi", level=1)
    h.runs[0].font.color.rgb = RGBColor(25, 118, 210)

    doc.add_paragraph(
        "Uygulamayı çalıştırmak için terminalde şu komut yeterli:"
    )

    p = doc.add_paragraph("python3 main.py")
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        run.bold = True

    doc.add_paragraph(
        "Pencere açıldığında üst kısımda üç mod butonu görünür: "
        "Seri/Paralel, Hata Ağacı ve Olay Ağacı. İstenen moda tıklayarak "
        "geçiş yapılır. Sol panelde veri girişi, sağ panelde diyagram "
        "ve sonuçlar yer alır."
    )

    # FTA kullanım
    p = doc.add_paragraph()
    run = p.add_run("Hata Ağacı modu: ")
    run.bold = True
    p.add_run(
        "Önce \"Olay Ekle\" butonuyla temel olayları tanımlayın (isim ve "
        "güvenilirlik değeri girin). Sonra \"Kapı Ekle\" ile mantık geçitlerini "
        "oluşturun — geçit tipini (AND/OR/VOTE) seçip alt elemanlarını "
        "işaretleyin. En üstteki geçidi \"Tepe Olay Seç\" ile belirleyin. "
        "Sonuçlar ve diyagram otomatik güncellenir."
    )

    # ETA kullanım
    p = doc.add_paragraph()
    run = p.add_run("Olay Ağacı modu: ")
    run.bold = True
    p.add_run(
        "Başlatıcı olay adını ve frekansını girin. \"Dal Ekle\" ile "
        "güvenlik bariyerlerini sırayla ekleyin (isim ve başarı olasılığı). "
        "Program tüm olası sonuçları otomatik hesaplar."
    )

    # SP kullanım
    p = doc.add_paragraph()
    run = p.add_run("Seri/Paralel modu: ")
    run.bold = True
    p.add_run(
        "Sol paneldeki ağaç yapısından düğüm seçip sağ tıklayarak bileşen "
        "veya grup (seri/paralel) ekleyin. Bileşenlerin güvenilirlik değerini "
        "girin, sistem güvenilirliği otomatik hesaplanır."
    )

    # Ortak
    p = doc.add_paragraph()
    run = p.add_run("Ortak özellikler: ")
    run.bold = True
    p.add_run(
        "Dosya menüsünden projeyi JSON olarak kaydedip yükleyebilirsiniz. "
        "\"PDF Rapor\" butonu ile sonuçları dışa aktarabilirsiniz. "
        "\"Diyagramı Kaydet\" ile diyagramı PNG olarak saklayabilirsiniz. "
        "Hazır örnekleri denemek için Örnekler menüsünü kullanın."
    )

    # ══════════════════════════════════════
    #  6. ÖRNEKLER
    # ══════════════════════════════════════

    h = doc.add_heading("6. Örnek Analizler", level=1)
    h.runs[0].font.color.rgb = RGBColor(25, 118, 210)

    doc.add_paragraph(
        "Uygulamanın üç modunu da gösteren birer örnek çalıştırdım."
    )

    # 6.1 FTA
    h2 = doc.add_heading("6.1. Hata Ağacı — Redundant Pompa Sistemi", level=2)
    h2.runs[0].font.color.rgb = RGBColor(230, 81, 0)

    doc.add_paragraph(
        "Bir soğutma sisteminde iki pompa paralel çalışıyor (yedekli). "
        "Bunların yanında bir vana ve bir destek sistemi (güç + kontrol) var."
    )

    p = doc.add_paragraph()
    run = p.add_run("Ağaç yapısı:")
    run.bold = True

    tree_lines = [
        "Sistem Arızası (OR) ← TEPE OLAY",
        "  ├── Soğutma Kaybı (OR)",
        "  │     ├── Tüm Pompalar Arızalı (AND)",
        "  │     │     ├── Pompa A   R=0.99",
        "  │     │     └── Pompa B   R=0.98",
        "  │     └── Vana Arızası    R=0.95",
        "  └── Destek Sistemi Arızası (OR)",
        "        ├── Güç Kaybı       R=0.999",
        "        └── Kontrol Hatası  R=0.997",
    ]
    for line in tree_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Hesaplama:")
    run.bold = True

    calc_steps = [
        "AND geçidi (Pompalar): İkisi de bozulmalı → R = 1 − (0.01 × 0.02) = 0.9998",
        "OR geçidi (Soğutma): Biri yeter → R = 0.9998 × 0.95 = 0.9498",
        "OR geçidi (Destek): R = 0.999 × 0.997 = 0.9960",
        "OR geçidi (Tepe): R = 0.9498 × 0.9960 ≈ 0.946 → F ≈ 0.054",
    ]
    for step in calc_steps:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("MCS sonuçları: ")
    run.bold = True
    p.add_run(
        "Dört MCS bulundu. Vana tek başına bir MCS (1. derece) ve FV değeri "
        "en yüksek bileşen (~%92). Yani sistemi güçlendirmek istiyorsak "
        "önce vanayı iyileştirmeliyiz. AND geçidi sayesinde pompa grubu zaten "
        "çok güvenilir."
    )

    add_img(doc, "01_hata_agaci.png",
            "Şekil 1: Hata Ağacı Analizi — Redundant Pompa Sistemi")

    # 6.2 ETA
    h2 = doc.add_heading("6.2. Olay Ağacı — Büyük LOCA Senaryosu", level=2)
    h2.runs[0].font.color.rgb = RGBColor(230, 81, 0)

    doc.add_paragraph(
        "Nükleer reaktörde büyük bir boru kırılması (LOCA) olduğunu "
        "düşünelim. Başlatıcı olay frekansı 10⁻⁴/yıl. Sırayla 4 güvenlik "
        "bariyeri devreye giriyor:"
    )

    et_table = doc.add_table(rows=5, cols=3)
    et_table.style = 'Table Grid'
    et_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    et_headers = ["Bariyer", "P(başarı)", "Ne yapıyor?"]
    for i, h_text in enumerate(et_headers):
        cell = et_table.cell(0, i)
        cell.paragraphs[0].add_run(h_text).bold = True
        set_cell_shading(cell, "E65100")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    et_data = [
        ("Reaktör Trip", "0.999", "Kontrol çubukları düşüp reaksiyonu durdurur"),
        ("ECCS Enjeksiyonu", "0.99", "Acil soğutma sistemi devreye girer"),
        ("Konteyment Yalıtımı", "0.995", "Radyoaktif sızıntıyı engeller"),
        ("Uzun Vadeli Soğutma", "0.98", "Saatler boyunca artık ısıyı atar"),
    ]
    for row_idx, (name, prob, desc) in enumerate(et_data, 1):
        et_table.cell(row_idx, 0).text = name
        et_table.cell(row_idx, 1).text = prob
        et_table.cell(row_idx, 2).text = desc

    doc.add_paragraph()

    doc.add_paragraph(
        "4 bariyer → 2⁴ = 16 olası sonuç. Hepsi başarılı olursa P ≈ 0.964 "
        "(güvenli duruş). Hepsi başarısız olursa frekans ≈ 10⁻¹³/yıl — "
        "savunma derinliği sayesinde inanılmaz düşük bir değer."
    )

    add_img(doc, "02_olay_agaci.png",
            "Şekil 2: Olay Ağacı Analizi — Büyük LOCA Senaryosu")

    # 6.3 S/P
    h2 = doc.add_heading("6.3. Seri/Paralel — Güç Kaynağı Sistemi", level=2)
    h2.runs[0].font.color.rgb = RGBColor(230, 81, 0)

    doc.add_paragraph("Bir UPS sistemi dört bloktan oluşuyor:")

    sp_lines = [
        "Giriş → [Trafo] → [Redresör 1 ║ Redresör 2] → [Akü A ║ Akü B] → [İnvertör] → Çıkış",
        "         R=0.995     Paralel                    Paralel            R=0.98",
    ]
    for line in sp_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)

    doc.add_paragraph()

    sp_calc = [
        "Redresör grubu (paralel): R = 1 − (0.03)² = 0.9991",
        "Akü grubu (paralel): R = 1 − (0.01)² = 0.9999",
        "Tüm sistem (seri): R = 0.995 × 0.9991 × 0.9999 × 0.98 ≈ 0.974",
    ]
    for step in sp_calc:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        "Paralel yedeklilik ne kadar işe yaradığını burada net görüyoruz: "
        "tek redresör R=0.97 iken çift redresör R=0.9991'e çıkıyor. "
        "Zincirin zayıf halkası ise seri bağlı tek bileşenler (Trafo ve İnvertör)."
    )

    add_img(doc, "03_seri_paralel.png",
            "Şekil 3: Seri/Paralel Analizi — Güç Kaynağı Sistemi")

    # ══════════════════════════════════════
    #  7. DOĞRULAMA
    # ══════════════════════════════════════

    h = doc.add_heading("7. Doğrulama", level=1)
    h.runs[0].font.color.rgb = RGBColor(25, 118, 210)

    doc.add_paragraph(
        "Uygulamanın doğru hesaplayıp hesaplamadığını kontrol etmek için "
        "örnek sonuçları el hesabıyla karşılaştırdım."
    )

    ver_table = doc.add_table(rows=5, cols=4)
    ver_table.style = 'Table Grid'
    ver_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    ver_headers = ["Düğüm", "El Hesabı (R)", "Uygulama (R)", "Fark"]
    for i, h_text in enumerate(ver_headers):
        cell = ver_table.cell(0, i)
        cell.paragraphs[0].add_run(h_text).bold = True
        set_cell_shading(cell, "1976D2")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    ver_data = [
        ("AND (Pompalar)", "1 − 0.01×0.02 = 0.9998", "0.9998", "0"),
        ("OR (Soğutma)", "0.9998 × 0.95 = 0.9498", "0.9498", "0"),
        ("OR (Destek)", "0.999 × 0.997 = 0.9960", "0.9960", "0"),
        ("OR (Tepe)", "0.9498 × 0.9960 = 0.9460", "0.9460", "0"),
    ]
    for row_idx, (node, manual, app, diff) in enumerate(ver_data, 1):
        ver_table.cell(row_idx, 0).text = node
        ver_table.cell(row_idx, 1).text = manual
        ver_table.cell(row_idx, 2).text = app
        ver_table.cell(row_idx, 3).text = diff
        if row_idx % 2 == 0:
            for ci in range(4):
                set_cell_shading(ver_table.cell(row_idx, ci), "E8F5E9")

    doc.add_paragraph()
    doc.add_paragraph(
        "El hesabı ve uygulama sonuçları birebir aynı çıkıyor. "
        "Olay ağacı ve seri/paralel hesapları da aynı şekilde doğrulandı."
    )

    # ══════════════════════════════════════
    #  8. SONUÇ
    # ══════════════════════════════════════

    h = doc.add_heading("8. Sonuç", level=1)
    h.runs[0].font.color.rgb = RGBColor(25, 118, 210)

    doc.add_paragraph(
        "Bu projede derste öğrendiğimiz üç temel analiz yöntemini — FTA, ETA "
        "ve Seri/Paralel — birleştiren bir masaüstü uygulaması geliştirdim. "
        "Uygulama yaklaşık 4000 satır Python kodundan oluşuyor."
    )

    doc.add_paragraph(
        "Projeyi geliştirirken en çok dikkatimi çeken nokta, yedekliliğin "
        "güvenilirliği ne kadar artırdığı oldu. Örnek analizde tek bir pompa "
        "R = 0.99 iken iki pompayı AND geçidiyle (paralel) bağladığımızda "
        "R = 0.9998'e çıkıyor — arıza olasılığı 50 kat düşüyor. "
        "Ama aynı zamanda OR geçidindeki tek bir zayıf bileşen "
        "(vana, R = 0.95) tüm sistemi aşağı çekebiliyor. Bu durum "
        "\"zincir en zayıf halkası kadar güçlüdür\" ilkesini çok net gösteriyor."
    )

    doc.add_paragraph(
        "Önem ölçüleri de bu gözlemi destekliyor. Fussell-Vesely sonuçlarına "
        "göre vana toplam riskin ~%92'sinden sorumlu. Yani mühendislik "
        "kararı olarak pompaya değil vanaya yatırım yapmak çok daha etkili. "
        "Bu tür analizler olmadan sadece sezgiyle karar vermek yanıltıcı "
        "olabilir — çünkü en pahalı veya en büyük bileşen her zaman "
        "en kritik bileşen değildir."
    )

    doc.add_paragraph(
        "Olay ağacı tarafında ise savunma derinliğinin gücü dikkat çekici. "
        "LOCA senaryosunda her bir bariyer tek başına mükemmel değil "
        "(en düşüğü 0.98), ama dört katman üst üste geldiğinde en kötü "
        "senaryo frekansı 10⁻¹³/yıl'a düşüyor. Bu, nükleer güvenlik "
        "felsefesinin neden tek bir bariyere değil birden çok bariyere "
        "dayandığını sayısal olarak kanıtlıyor."
    )

    doc.add_paragraph(
        "Sonuç olarak bu proje, derste öğrendiğimiz kavramları sadece "
        "formül olarak değil, gerçek sistemler üzerinde uygulayarak "
        "anlamama yardımcı oldu. Bir sistemi modelleyip sonuçları görmek, "
        "\"R nedir, F nedir\" sorusundan \"bu sistemi nasıl iyileştiririm\" "
        "sorusuna geçmemi sağladı."
    )

    # ══════════════════════════════════════
    #  KAYNAKLAR
    # ══════════════════════════════════════

    doc.add_page_break()

    h = doc.add_heading("Kaynaklar", level=1)
    h.runs[0].font.color.rgb = RGBColor(25, 118, 210)

    references = [
        '[1] W. Vesely ve diğerleri, "Fault Tree Handbook," '
        'NUREG-0492, U.S. Nuclear Regulatory Commission, 1981.',

        '[2] U.S. NRC, "Reactor Safety Study," WASH-1400, 1975.',

        '[3] IEC 61025:2006, "Fault tree analysis (FTA)."',

        '[4] IEC 62502:2010, "Event tree analysis (ETA)."',

        '[5] M. Rausand, A. Hoyland, "System Reliability Theory," '
        '2nd ed., Wiley, 2004.',

        '[6] PyQt6 — https://www.riverbankcomputing.com/static/Docs/PyQt6/',

        '[7] Graphviz — https://graphviz.org/',
    ]

    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)

    # ══════════════════════════════════════
    #  SAYFA NUMARALARI
    # ══════════════════════════════════════

    add_page_numbers(doc)

    # ══════════════════════════════════════
    #  KAYDET
    # ══════════════════════════════════════

    output_path = os.path.join(os.path.dirname(__file__), "HÖA_Analizörü_Rapor.docx")
    doc.save(output_path)
    print(f"Rapor oluşturuldu: {output_path}")
    return output_path


if __name__ == "__main__":
    create_report()
